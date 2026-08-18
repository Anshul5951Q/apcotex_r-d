"""
app/services/pipeline/report_service.py

Aggregates extracted patent data, generates a Markdown report via Gemini,
and exports the final report to PDF and DOCX formats.
"""
import logging
import os
import uuid
from typing import List, Tuple, Dict, Optional

import markdown
from docx import Document
from xhtml2pdf import pisa

from app.services.llm.llm_client import llm_client
from app.services.pipeline.schemas import (
    ReportPatentEvidence, PatentResearchReport, LLMPatentResearchReport,
    ReportPatent, ReportPatentDetails, ReportPatentMethodology, PatentExtraction
)
from app.services.prompts.patent_prompts import (
    REPORT_GENERATION_SYSTEM_PROMPT,
    REPORT_GENERATION_USER_TEMPLATE
)

logger = logging.getLogger(__name__)

class ReportService:
    def __init__(self):
        # Ensure export directory exists
        self.export_dir = os.path.join(os.getcwd(), "exports")
        os.makedirs(self.export_dir, exist_ok=True)

    async def generate_structured_report(
        self, compound_name: str, extractions: List[ReportPatentEvidence],
        patent_manifest: List[str] = None,
        secondary_candidates: list = None  # list[SearchResult] or similar objects
    ) -> tuple:
        """Generate the structured report via LLM using the aggregated extractions."""
        import time
        from app.core.config import settings

        logger.info("Generating final structured report for %d patents...", len(extractions))

        from app.services.pipeline.report_evidence_service import ReportEvidenceService
        svc = ReportEvidenceService()

        effective_provider_limit = getattr(settings, 'REPORT_PROVIDER_SAFE_LIMIT', 100000)
        overhead = getattr(settings, 'REPORT_EVIDENCE_OVERHEAD_TOKENS', 4000)
        evidence_budget = max(1000, effective_provider_limit - overhead)

        sys_prompt = REPORT_GENERATION_SYSTEM_PROMPT.format(compound_name=compound_name)
        sys_tokens = svc.estimate_tokens(sys_prompt)
        # Estimate base overhead using dummy values for all template placeholders
        base_tokens = svc.estimate_tokens(
            REPORT_GENERATION_USER_TEMPLATE.format(
                compound_name=compound_name,
                extractions_data="",
                patent_manifest="",
                secondary_manifest="None",
                primary_count=0,
                patent_count=0
            )
        )
        overhead_total = sys_tokens + base_tokens


        # ── Safety: trim evidence per-patent proportionally if still oversized ────
        extractions_data = svc.serialize_evidence(extractions)
        est_tokens = svc.estimate_tokens(extractions_data)
        total_prompt_tokens = est_tokens + overhead_total

        if total_prompt_tokens > effective_provider_limit:
            deficit = total_prompt_tokens - effective_provider_limit
            logger.warning(
                "[REPORT] Evidence still oversized after upstream compaction: "
                "%d tokens (limit %d, deficit %d). Applying per-patent proportional trim.",
                total_prompt_tokens, effective_provider_limit, deficit
            )
            # Trim source sentences proportionally across all patents (never drop a patent)
            for ev in extractions:
                for param in ev.overall_patent_parameters:
                    if param.source_sentence and len(param.source_sentence) > 60:
                        param.source_sentence = param.source_sentence[:60] + "..."
                for ex in ev.examples:
                    for param in ex.extracted_parameters:
                        if param.source_sentence and len(param.source_sentence) > 60:
                            param.source_sentence = param.source_sentence[:60] + "..."
            extractions_data = svc.serialize_evidence(extractions)
            est_tokens = svc.estimate_tokens(extractions_data)
            total_prompt_tokens = est_tokens + overhead_total
            logger.info(
                "[REPORT] After proportional trim: %d tokens (limit %d)",
                total_prompt_tokens, effective_provider_limit
            )

        # Build patent manifest for injection into user prompt
        if patent_manifest is None:
            patent_manifest = [ev.patent_number for ev in extractions]
        manifest_lines = [f"{i+1}. {pn}" for i, pn in enumerate(patent_manifest)]
        manifest_str = "\n".join(manifest_lines)
        patent_count = len(patent_manifest)
        primary_count = patent_count

        prompt = REPORT_GENERATION_USER_TEMPLATE.format(
            compound_name=compound_name,
            extractions_data=extractions_data,
            patent_manifest=manifest_str,
            primary_count=primary_count,
            patent_count=patent_count
        )

        logger.info(
            "REPORT PAYLOAD\nPatents: %d\nEvidence tokens: %d\nSystem tokens: %d\nTotal input tokens: %d\nConfigured limit: %d",
            len(extractions), est_tokens, overhead_total, total_prompt_tokens, effective_provider_limit
        )

        t0 = time.time()
        try:
            report_obj, provider_id, _usage = await llm_client.generate_structured(
                prompt=prompt,
                system_prompt=sys_prompt,
                schema=LLMPatentResearchReport,
                temperature=0.2
            )
            latency = time.time() - t0
            in_tokens = (_usage or {}).get("input_tokens", est_tokens)
            out_tokens = (_usage or {}).get("output_tokens", 0)

            if report_obj is None:
                logger.error(
                    "[REPORT] RESPONSE_EMPTY | Latency: %.1fs | Provider: %s | "
                    "Input tokens: %d | Status: FAILED",
                    latency, provider_id, in_tokens
                )
                raise ValueError("Report generation failed: LLM returned None instead of structured object")

            # Deterministically map extracted patent evidence to the final report
            methodology_patents = []
            for ext in extractions:
                details = ReportPatentDetails(
                    patent_number=ext.patent_number,
                    patent_title=ext.title,
                    assignee=ext.assignee,
                    publication_year=ext.publication_year,
                    jurisdiction=ext.jurisdiction,
                    legal_status="Unknown",
                    polymer_type="Not disclosed",
                    relevance_to_target="Automatically extracted candidate",
                    relevance_tier="PRIMARY"
                )
                
                params = []
                for p in ext.overall_patent_parameters:
                    s = f"{p.name}: {p.value} {p.unit}".strip()
                    if p.context:
                        s += f" ({p.context})"
                    params.append(s)
                    
                methodology = ReportPatentMethodology(dynamic_parameters=params)
                
                evidence = []
                for findings in ext.technical_findings:
                    evidence.append(findings)
                for ex in ext.examples:
                    evidence.append(f"Example {ex.example_id}: " + ", ".join([f"{p.name}: {p.value} {p.unit}" for p in ex.extracted_parameters]))
                    
                methodology_patents.append(ReportPatent(
                    patent_details=details,
                    polymerization_method=methodology,
                    experimental_evidence=evidence if evidence else ["No specific experimental examples disclosed."],
                    technical_relevance="Selected via deterministic pipeline scoring."
                ))
            
            final_report = PatentResearchReport(
                title=report_obj.title or "PATENT RESEARCH REPORT",
                abstract=report_obj.abstract or "No abstract provided.",
                methodology_patents=methodology_patents,
                cross_patent_comparison=report_obj.cross_patent_comparison,
                conclusion=report_obj.conclusion,
                references=report_obj.references
            )

            logger.info(
                "[REPORT] RESPONSE_RECEIVED | Latency: %.1fs | Provider: %s | "
                "Input tokens: %d | Output tokens: %d | Patents mapped: %d | Status: SUCCESS",
                latency, provider_id, in_tokens, out_tokens,
                len(final_report.methodology_patents)
            )

            return final_report, _usage

        except Exception as e:
            latency = time.time() - t0
            logger.error(
                "[REPORT] REPORT_FAILED | Type: %s | Message: %s | "
                "Latency: %.1fs | Provider limit: %d | Input tokens estimated: %d",
                type(e).__name__, str(e)[:300], latency, effective_provider_limit, total_prompt_tokens
            )
            raise


    def report_to_markdown(self, report: 'PatentResearchReport') -> str:
        """Converts the structured report to markdown for PDF export (canonical data)."""
        lines = []

        title = report.title or "PATENT RESEARCH REPORT"
        lines.append(f"# {title.upper()}")

        abstract = report.abstract or "No abstract provided."
        lines.append("\n## 1. ABSTRACT")
        lines.append(abstract)

        lines.append("\n## 2. METHODOLOGY")
        lines.append("### PRIMARY PATENT EVIDENCE")

        primary_patents = getattr(report, 'methodology_patents', [])
        secondary_patents = getattr(report, 'secondary_patents', [])

        for idx, patent in enumerate(primary_patents):
            lines.append(f"\n#### Patent {idx + 1}")
            pd = patent.patent_details
            lines.append(f"- Patent Number: {pd.patent_number or 'Not available from source'}")
            lines.append(f"- Patent Title: {pd.patent_title or 'Not available from source'}")
            lines.append(f"- Assignee: {pd.assignee or 'Not available from source'}")
            lines.append(f"- Jurisdiction: {pd.jurisdiction or 'Not available from source'}")
            lines.append(f"- Publication Year: {pd.publication_year or 'Not available from source'}")
            if getattr(pd, 'priority_date', None):
                lines.append(f"- Priority Date: {pd.priority_date}")
            if getattr(pd, 'legal_status', None):
                lines.append(f"- Legal Status: {pd.legal_status}")
            if getattr(pd, 'polymer_type', None):
                lines.append(f"- Polymer Type: {pd.polymer_type}")
            lines.append(f"- Relevance: {pd.relevance_to_target or 'Not disclosed'}")

            method = patent.polymerization_method
            lines.append("\n**Polymerization / Synthesis Method**")
            if method and method.dynamic_parameters:
                for param_str in method.dynamic_parameters:
                    lines.append(f"- {param_str}")
            else:
                lines.append("- No polymerization parameters disclosed.")

            lines.append("\n**Relevant Experimental Evidence**")
            if patent.experimental_evidence:
                for ev in patent.experimental_evidence:
                    lines.append(f"- {ev}")
            else:
                lines.append("- No experimental evidence provided.")

            lines.append("\n**Technical Relevance**")
            lines.append(patent.technical_relevance or "No technical relevance provided.")

        if secondary_patents:
            lines.append("\n### SUPPORTING / RELATED PATENTS (SECONDARY)")
            for idx, patent in enumerate(secondary_patents):
                lines.append(f"\n#### Supporting Patent {idx + 1}")
                pd = patent.patent_details
                lines.append(f"- Patent Number: {pd.patent_number or 'Not available from source'}")
                lines.append(f"- Patent Title: {pd.patent_title or 'Not available from source'}")
                lines.append(f"- Assignee: {pd.assignee or 'Not available from source'}")
                lines.append(f"- Jurisdiction: {pd.jurisdiction or 'Not available from source'}")
                lines.append(f"- Classification: SECONDARY (Synthesis relevant; constraints not confirmed)")
                lines.append(f"- Technical Note: {patent.technical_relevance or 'Supporting patent.'}")

        # Cross-patent comparison — only when enough primary evidence
        lines.append("\n## 3. CROSS-PATENT COMPARISON & SYNTHESIS TRENDS")
        if len(primary_patents) >= 2 and report.cross_patent_comparison:
            for point in report.cross_patent_comparison:
                lines.append(f"- {point}")
        elif len(primary_patents) < 2:
            lines.append(
                "- No qualifying primary patents were identified under the configured "
                "target material relevance criteria. Cross-patent quantitative "
                "trends were therefore not generated."
            )
        else:
            lines.append("- No cross-patent comparison data provided.")

        # Conclusion
        if report.conclusion:
            lines.append("\n## 4. CONCLUSION")
            lines.append(report.conclusion)
            next_section = 5
        else:
            next_section = 4

        # References — from validated evidence set only
        lines.append(f"\n## {next_section}. REFERENCES")
        if report.references:
            for ref in report.references:
                lines.append(f"- {ref}")
        else:
            lines.append("- No references provided.")

        return "\n".join(lines)

    async def export_to_pdf(self, markdown_text: str, file_name: str) -> str:
        """Convert Markdown to HTML, then to PDF using xhtml2pdf."""
        logger.info("Exporting report to PDF: %s", file_name)
        
        # Convert Markdown to HTML
        html_content = markdown.markdown(markdown_text, extensions=['tables'])
        
        # Basic APCOTEX CSS styling wrapper
        styled_html = f"""
        <html>
        <head>
            <style>
                @page {{ size: a4 portrait; margin: 2cm; }}
                body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.5; color: #333; }}
                h1 {{ color: #004080; font-size: 24pt; border-bottom: 2px solid #004080; padding-bottom: 5px; }}
                h2 {{ color: #0059b3; font-size: 18pt; margin-top: 20px; }}
                h3 {{ color: #0073e6; font-size: 14pt; margin-top: 15px; }}
                table {{ width: 100%; border-collapse: collapse; margin-top: 10px; margin-bottom: 10px; }}
                th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        pdf_path = os.path.join(self.export_dir, file_name)
        
        try:
            with open(pdf_path, "w+b") as result_file:
                pisa_status = pisa.CreatePDF(styled_html, dest=result_file)
                
            if pisa_status.err:
                logger.error("Error creating PDF via xhtml2pdf")
                raise Exception("PDF generation failed.")
                
            return pdf_path
        except Exception as e:
            logger.error("Export to PDF failed: %s", e)
            return ""

    def _sanitize_text_for_xml(self, text: str) -> str:
        """Remove control characters and NULL bytes that are incompatible with XML."""
        if not text:
            return text
        # Remove control characters except tab, newline, carriage return
        # XML 1.0 allows: #x9, #xA, #xD
        import re
        # Remove all control characters except tab, newline, carriage return
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        return text

    async def export_to_docx(self, report: 'PatentResearchReport', file_name: str) -> str:
        """
        Export the DOCX directly from the canonical PatentResearchReport structured object.
        This ensures zero data loss between web report and DOCX.
        Signature changed: accepts PatentResearchReport, not markdown_text.
        """
        logger.info("Exporting report to DOCX (structured): %s", file_name)

        doc = Document()

        def _safe(text) -> str:
            """Return text as a clean string, falling back to placeholder."""
            if text is None:
                return "Not available from source"
            s = str(text).strip()
            if not s or s.lower() in ("not disclosed", "none"):
                return "Not available from source"
            # Remove control characters incompatible with XML
            import re as _re
            return _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)

        def _add_metadata_table(doc, items: list[tuple[str, str]]):
            """Add a two-column metadata table (label, value)."""
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            for label, value in items:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = _safe(value)
                # Bold the label
                for paragraph in row[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Title
        doc.add_heading(_safe(report.title or 'PATENT RESEARCH REPORT'), level=1)

        # Abstract
        doc.add_heading('1. Abstract', level=2)
        doc.add_paragraph(_safe(report.abstract))

        # Methodology / Primary Patents
        primary_patents = getattr(report, 'methodology_patents', [])
        secondary_patents = getattr(report, 'secondary_patents', [])

        doc.add_heading('2. Primary Patent Evidence', level=2)

        if not primary_patents:
            doc.add_paragraph(
                'No qualifying primary patents were identified under the configured '
                'target material relevance criteria.'
            )
        else:
            for idx, patent in enumerate(primary_patents):
                pd = patent.patent_details
                doc.add_heading(f'Patent {idx + 1}: {_safe(pd.patent_number)} — {_safe(pd.patent_title)}', level=3)

                # Metadata table
                metadata_items = [
                    ('Patent Number', pd.patent_number),
                    ('Title', pd.patent_title),
                    ('Assignee', pd.assignee),
                    ('Jurisdiction', pd.jurisdiction),
                    ('Publication Year', pd.publication_year),
                    ('Priority Date', getattr(pd, 'priority_date', None)),
                    ('Legal Status', getattr(pd, 'legal_status', None)),
                    ('Polymer Type', getattr(pd, 'polymer_type', None)),
                    ('Relevance Classification', 'PRIMARY'),
                    ('Relevance to Target', pd.relevance_to_target),
                ]
                _add_metadata_table(doc, metadata_items)

                # Polymerization parameters
                doc.add_heading('Polymerization / Synthesis Method', level=4)
                method = patent.polymerization_method
                if method and method.dynamic_parameters:
                    for param_str in method.dynamic_parameters:
                        doc.add_paragraph(_safe(param_str), style='List Bullet')
                else:
                    doc.add_paragraph('No polymerization parameters disclosed.')

                # Experimental evidence
                doc.add_heading('Relevant Experimental Evidence', level=4)
                if patent.experimental_evidence:
                    for ev in patent.experimental_evidence:
                        doc.add_paragraph(_safe(ev), style='List Bullet')
                else:
                    doc.add_paragraph('No experimental evidence provided.')

                # Technical relevance
                doc.add_heading('Technical Relevance', level=4)
                doc.add_paragraph(_safe(patent.technical_relevance))

        # Secondary patents section
        if secondary_patents:
            doc.add_heading('Supporting / Related Patents (Secondary)', level=2)
            doc.add_paragraph(
                'These patents are related to target synthesis but specific constraint relevance was not '
                'confirmed from available title and abstract data.'
            )
            for idx, patent in enumerate(secondary_patents):
                pd = patent.patent_details
                doc.add_heading(f'Supporting Patent {idx + 1}: {_safe(pd.patent_number)} — {_safe(pd.patent_title)}', level=3)
                meta = [
                    ('Patent Number', pd.patent_number),
                    ('Title', pd.patent_title),
                    ('Assignee', pd.assignee),
                    ('Jurisdiction', pd.jurisdiction),
                    ('Publication Year', pd.publication_year),
                    ('Relevance Classification', 'SECONDARY'),
                ]
                _add_metadata_table(doc, meta)
                doc.add_heading('Technical Note', level=4)
                doc.add_paragraph(_safe(patent.technical_relevance))

        # Cross-patent comparison
        doc.add_heading('3. Cross-Patent Comparison & Synthesis Trends', level=2)
        if len(primary_patents) >= 2 and report.cross_patent_comparison:
            for point in report.cross_patent_comparison:
                doc.add_paragraph(_safe(point), style='List Bullet')
        elif len(primary_patents) < 2:
            doc.add_paragraph(
                'No qualifying primary patents were identified under the configured '
                'target material relevance criteria. Cross-patent quantitative '
                'trends were therefore not generated.'
            )
        else:
            doc.add_paragraph('Insufficient comparable evidence for cross-patent trend analysis.')

        # Conclusion
        conclusion = getattr(report, 'conclusion', None)
        if conclusion:
            doc.add_heading('4. Conclusion', level=2)
            doc.add_paragraph(_safe(conclusion))
            ref_section_num = 5
        else:
            ref_section_num = 4

        # References
        doc.add_heading(f'{ref_section_num}. References', level=2)
        if report.references:
            for ref in report.references:
                doc.add_paragraph(_safe(ref), style='List Bullet')
        else:
            doc.add_paragraph('No references provided.')

        docx_path = os.path.join(self.export_dir, file_name)
        try:
            doc.save(docx_path)
            return docx_path
        except Exception as e:
            logger.error("Export to DOCX failed: %s", e)
            return ""

    def validate_report_consistency(
        self,
        report: 'PatentResearchReport',
        primary_manifest: list[str],
        secondary_manifest: list[str] = None
    ) -> tuple[bool, list[str]]:
        """
        Run deterministic consistency checks before saving a report run.

        Returns:
            (ok, errors) where ok=True means all checks passed.
        """
        errors = []
        secondary_manifest = secondary_manifest or []
        all_manifest = set(primary_manifest) | set(secondary_manifest)

        primary_patents = getattr(report, 'methodology_patents', [])
        secondary_patents = getattr(report, 'secondary_patents', [])

        # Check 1: Every primary patent has a publication_number
        for p in primary_patents:
            pn = getattr(p.patent_details, 'patent_number', None)
            if not pn:
                errors.append(f"CHECK FAIL: Primary patent missing patent_number: {p}")

        # Check 2: Abstract exists
        if not report.abstract:
            errors.append("CHECK FAIL: Report abstract is empty.")

        # Check 3: If 0 primary patents, cross_patent_comparison must be empty
        if len(primary_patents) == 0 and report.cross_patent_comparison:
            errors.append(
                "CHECK FAIL: 0 primary patents but cross_patent_comparison is non-empty."
            )

        # Check 4: If < 2 primary patents, cross_patent_comparison must be empty
        if len(primary_patents) < 2 and report.cross_patent_comparison:
            errors.append(
                f"CHECK FAIL: {len(primary_patents)} primary patent(s) but "
                "cross_patent_comparison is non-empty (requires >= 2)."
            )

        # Check 5: Every reference must be from the manifest
        for ref in report.references:
            # References format: 'PatentNumber | ...'
            ref_pn = ref.split('|')[0].strip() if '|' in ref else ref.strip()
            if ref_pn and all_manifest and ref_pn not in all_manifest:
                errors.append(f"CHECK FAIL: Reference '{ref_pn}' not in evidence manifest.")

        # Check 6: Primary patent numbers match manifest
        report_primary_pns = {
            getattr(p.patent_details, 'patent_number', '') for p in primary_patents
        }
        manifest_set = set(primary_manifest)
        missing = manifest_set - report_primary_pns
        if missing:
            errors.append(f"CHECK FAIL: Missing patents in report: {sorted(missing)}")

        ok = len(errors) == 0
        if ok:
            logger.info("[REPORT CONSISTENCY] All checks PASSED.")
        else:
            for err in errors:
                logger.error("[REPORT CONSISTENCY] %s", err)
        return ok, errors
